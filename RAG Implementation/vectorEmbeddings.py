"""Here we are going through the list of files (Research paper summaries) and extracting the metadata from each file. 
After extracting the metadata from each of the files we are going to create a dictionary with the file name as the key and the metadata as the value.
Then we go ahead and create chunks for the rest of the data in the respective files. We print the metadata and chunks for each file.
We then create vector embeddings for each chunk and send it to the database."""

import os
import re
import torch.nn.functional as F
import torch
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
from transformers import AutoTokenizer, AutoModel


folder_path = "C:\\Users\\kshit\\Jupyteranaconda\\summaries"

metadata = {}

def extract_metadata(doc_text):
    title = re.search(r"Title:\s*(.*)", doc_text, re.IGNORECASE)
    authors = re.search(r"Authors?:\s*(.*)", doc_text, re.IGNORECASE)
    doi = re.search(r"DOI:\s*(.*)", doc_text, re.IGNORECASE)
    
    return {
        "Title": title.group(1).strip() if title else None,
        "Authors": authors.group(1).strip() if authors else None,
        "DOI": doi.group(1).strip() if doi else None,
    }

for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):  
        file_path = os.path.join(folder_path, filename)
   
        doc = Document(file_path)

        doc_text = "\n".join([para.text for para in doc.paragraphs])

        metadata[filename] = extract_metadata(doc_text)

for filename, details in metadata.items():
    print(f"File: {filename}")
    print(f"Title: {details['Title']}")
    print(f"Authors: {details['Authors']}")
    print(f"DOI: {details['DOI']}")
    print("-" * 40)  


def remove_metadata(doc_text):
    doc_text = re.sub(r"(Title:\s*.*\n|Authors?:\s*.*\n|DOI:\s*.*\n)", "", doc_text)
    return doc_text.strip()

for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        
        doc = Document(file_path)
        
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        
        metadata[filename] = extract_metadata(doc_text)
        content_text = remove_metadata(doc_text)
        splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
        chunks = splitter.split_text(content_text)
        metadata[filename]["Chunks"] = chunks

for filename, details in metadata.items():
    print(f"File: {filename}")
    print("Chunks:")
    for i, chunk in enumerate(details["Chunks"]):
        print(f"Chunk {i+1}: {chunk}")
        print("-" * 20)
    print("-" * 40)  
    

for filename, details in metadata.items():
    print(f"File: {filename}")
    print(f"Title: {details['Title']}")
    print(f"Authors: {details['Authors']}")
    print(f"DOI: {details['DOI']}")
    print("Chunks:")
    for i, chunk in enumerate(details["Chunks"]):
        print(f"Chunk {i+1}: {chunk}")
        print("-" * 20)
    print("-" * 40)
    

model_ckpt = "sentence-transformers/all-MiniLM-L6-v2"
tokenizer = AutoTokenizer.from_pretrained(model_ckpt)
model = AutoModel.from_pretrained(model_ckpt)

for filename, details in metadata.items():
    chunks = details["Chunks"]
    encoded_input = tokenizer(chunks, padding=True, truncation=True, return_tensors="pt")
    
    with torch.no_grad():
        model_output = model(**encoded_input)
    
    def mean_pooling(model_output, attention_mask):
        token_embeddings = model_output.last_hidden_state
        input_mask_expanded = (
            attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
        )
        return torch.sum(token_embeddings * input_mask_expanded, 1) / torch.clamp(
            input_mask_expanded.sum(1), min=1e-9
        )

    sentence_embeddings = mean_pooling(model_output, encoded_input["attention_mask"])
    sentence_embeddings = F.normalize(sentence_embeddings, p=2, dim=1)
    sentence_embeddings_np = sentence_embeddings.detach().numpy()
    
    details["Embeddings"] = sentence_embeddings_np

for filename, details in metadata.items():
    print(f"File: {filename}")
    print(f"Title: {details['Title']}")
    print(f"Authors: {details['Authors']}")
    print(f"DOI: {details['DOI']}")
    print("Chunks and their embeddings:")
    for i, chunk in enumerate(details["Chunks"]):
        embedding = details["Embeddings"][i]
        print(f"Chunk {i+1}:")
        print(f"Text: {chunk}")
        print(f"Embedding {i+1}: {embedding}\n")
    print("-" * 40)  
